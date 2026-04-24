from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path("docs/ReconXpose_Internship_Report.md")
OUT = Path("docs/ReconXpose_Internship_Report.docx")


def set_margins(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def style_document(doc: Document):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    for style_name, size in [('Title', 16), ('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)]:
        style = styles[style_name]
        style.font.name = 'Times New Roman'
        style.font.size = Pt(size)
        style.font.bold = True


def parse_markdown(md_text: str):
    lines = md_text.splitlines()
    blocks = []
    buf = []
    current = ('p', '')
    for line in lines:
        if not line.strip():
            if buf:
                blocks.append((current[0], '\n'.join(buf)))
                buf = []
            continue
        if line.startswith('# '):
            if buf:
                blocks.append((current[0], '\n'.join(buf)))
                buf = []
            current = ('h1', line[2:].strip())
            blocks.append(current)
        elif line.startswith('## '):
            if buf:
                blocks.append((current[0], '\n'.join(buf)))
                buf = []
            current = ('h2', line[3:].strip())
            blocks.append(current)
        elif line.startswith('### '):
            if buf:
                blocks.append((current[0], '\n'.join(buf)))
                buf = []
            current = ('h3', line[4:].strip())
            blocks.append(current)
        else:
            buf.append(line)
    if buf:
        blocks.append((current[0], '\n'.join(buf)))
    return blocks


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def build_docx():
    text = SRC.read_text(encoding='utf-8')
    doc = Document()
    set_margins(doc.sections[0])
    style_document(doc)

    # header/footer
    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run('ReconXpose Internship Report').font.name = 'Times New Roman'

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    for kind, content in parse_markdown(text):
        if kind == 'h1':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(content)
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(16)
        elif kind == 'h2':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(content)
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
        elif kind == 'h3':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(content)
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
        else:
            if content.startswith('- '):
                for item in content.splitlines():
                    if item.startswith('- '):
                        add_paragraph(doc, f"• {item[2:]}")
            else:
                add_paragraph(doc, content)

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    build_docx()
